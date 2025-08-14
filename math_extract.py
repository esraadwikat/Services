import fitz  # PyMuPDF
from docx import Document
import re

# ===============================
# مسار الكتاب
# ===============================
pdf_path = r"C:\Users\Sabitu\Services_Folder\math.pdf"  # ضع مسار كتابك هنا
doc = fitz.open(pdf_path)

# إنشاء ملف Word جديد
word_doc = Document()
word_doc.add_heading("Extracted Math Book", 0)

# ===============================
# دالة لتنقية النصوص من أحرف غير صالحة
# ===============================
def clean_text(text):
    # إزالة أحرف التحكم (control characters) ما عدا السطر الجديد
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
    return text

# ===============================
# دالة لفصل الأسئلة إذا كانت مرقمة
# ===============================
def split_questions(text):
    # يفصل النص بعد أي رقم متبوع بنقطة (1. 2. 3. ...)
    parts = re.split(r'(\d+\.)', text)
    questions = []
    for i in range(1, len(parts), 2):
        number = parts[i]
        content = parts[i+1] if i+1 < len(parts) else ""
        questions.append(number + content.strip())
    return questions

# ===============================
# استخراج النصوص لكل صفحة
# ===============================
for page_num, page in enumerate(doc, start=1):
    word_doc.add_heading(f"Page {page_num}", level=1)
    
    # استخراج النصوص بالترتيب الطبيعي للصفحة
    text = page.get_text("text")
    text = clean_text(text)
    
    # تقسيم النص إلى أسئلة إذا مرقمة
    questions = split_questions(text)
    
    if questions:
        for q in questions:
            word_doc.add_paragraph(q)
    else:
        # إذا لم يتم العثور على أسئلة مرقمة، أضف النص كله
        word_doc.add_paragraph(text)

# ===============================
# حفظ الملف
# ===============================
output_path = "Extracted_Math_Book.docx"
word_doc.save(output_path)
print(f"Extraction complete! Saved as {output_path}")